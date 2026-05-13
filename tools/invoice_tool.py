import os
import io
import streamlit as st
import pandas as pd
import sqlite3
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .invoice_templates import (
    get_all_templates, get_default_template, save_template, 
    set_default_template, delete_template, create_blank_template
)

class InvoicePDF(FPDF):
    def __init__(self, template=None):
        super().__init__()
        self.template = template or get_default_template()
    
    def header(self):
        title = self.template.get("title", "Invoice")
        title_font = self.template.get("title_font", "Arial")
        title_size = self.template.get("title_size", 14)
        title_bold = self.template.get("title_bold", True)
        
        self.set_font(title_font, 'B' if title_bold else '', title_size)
        self.cell(0, 10, title, 0, 1, 'C')
        self.ln(10)

    def footer(self):
        company_footer = self.template.get("company_footer", "")
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, company_footer, 0, 0, 'C')

def generate_invoice_pdf(loan_data, client_data, payments_data, template):
    """Generate a customized PDF invoice based on template"""
    pdf = InvoicePDF(template)
    pdf.add_page()
    
    heading_font = template.get("heading_font", "Arial")
    heading_size = template.get("heading_size", 11)
    heading_bold = template.get("heading_bold", True)
    body_font = template.get("body_font", "Arial")
    body_size = template.get("body_size", 10)
    body_bold = template.get("body_bold", False)
    
    include_sections = template.get("include_sections", {})
    
    # CLIENT INFORMATION
    if include_sections.get("client_info", True):
        pdf.set_font(heading_font, 'B' if heading_bold else '', heading_size)
        pdf.cell(0, 10, 'Client Information:', 0, 1)
        pdf.set_font(body_font, 'B' if body_bold else '', body_size)
        
        client_fields = template.get("client_info_fields", {})
        
        if client_fields.get("first_name", True) and client_fields.get("last_name", True):
            pdf.cell(0, 10, f'Name: {client_data["first_name"]} {client_data["last_name"]}', 0, 1)
        
        if client_fields.get("id_number", True):
            pdf.cell(0, 10, f'ID Number: {client_data["id_number"]}', 0, 1)
        
        if client_fields.get("phone", True):
            pdf.cell(0, 10, f'Phone: {client_data["phone"]}', 0, 1)
        
        pdf.ln(5)
    
    # LOAN DETAILS
    if include_sections.get("loan_details", True):
        pdf.set_font(heading_font, 'B' if heading_bold else '', heading_size)
        pdf.cell(0, 10, 'Loan Details:', 0, 1)
        pdf.set_font(body_font, 'B' if body_bold else '', body_size)
        
        loan_fields = template.get("loan_details_fields", {})
        
        if loan_fields.get("loan_id", True):
            pdf.cell(0, 10, f'Loan ID: {loan_data["loan_id"]}', 0, 1)
        if loan_fields.get("principal", True):
            pdf.cell(0, 10, f'Principal Amount: R {loan_data["principal"]:.2f}', 0, 1)
        if loan_fields.get("balance", True):
            pdf.cell(0, 10, f'Current Balance: R {loan_data["balance"]:.2f}', 0, 1)
        if loan_fields.get("amount_paid", True):
            pdf.cell(0, 10, f'Amount Paid: R {loan_data["amount_paid"]:.2f}', 0, 1)
        if loan_fields.get("due_date", True):
            pdf.cell(0, 10, f'Due Date: {loan_data["due_date"]}', 0, 1)
        if loan_fields.get("status", True):
            pdf.cell(0, 10, f'Status: {loan_data["status"]}', 0, 1)
        
        pdf.ln(5)
    
    # PAYMENT HISTORY
    if include_sections.get("payment_history", True) and not payments_data.empty:
        pdf.set_font(heading_font, 'B' if heading_bold else '', heading_size)
        pdf.cell(0, 10, 'Payment History:', 0, 1)
        pdf.set_font(body_font, '', 8)
        
        payment_fields = template.get("payment_history_fields", {})
        col_width = 60
        
        if payment_fields.get("date", True):
            pdf.cell(col_width, 10, 'Date', 1)
        if payment_fields.get("amount", True):
            pdf.cell(col_width, 10, 'Amount', 1)
        if payment_fields.get("type", True):
            pdf.cell(col_width, 10, 'Type', 1)
        pdf.ln()
        
        for _, payment in payments_data.iterrows():
            if payment_fields.get("date", True):
                pdf.cell(col_width, 10, str(payment['date']), 1)
            if payment_fields.get("amount", True):
                pdf.cell(col_width, 10, f'R {payment["amount"]:.2f}', 1)
            if payment_fields.get("type", True):
                pdf.cell(col_width, 10, str(payment['type']), 1)
            pdf.ln()
    
    return pdf

def generate_invoice_docx(loan_data, client_data, payments_data, template):
    """Generate a customized Word document (fully editable)"""
    doc = Document()
    
    title = template.get("title", "Invoice")
    title_size = template.get("title_size", 14)
    
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    include_sections = template.get("include_sections", {})
    
    # CLIENT INFORMATION
    if include_sections.get("client_info", True):
        doc.add_heading('Client Information:', level=2)
        client_fields = template.get("client_info_fields", {})
        
        if client_fields.get("first_name", True) and client_fields.get("last_name", True):
            doc.add_paragraph(f'Name: {client_data["first_name"]} {client_data["last_name"]}')
        if client_fields.get("id_number", True):
            doc.add_paragraph(f'ID Number: {client_data["id_number"]}')
        if client_fields.get("phone", True):
            doc.add_paragraph(f'Phone: {client_data["phone"]}')
    
    # LOAN DETAILS
    if include_sections.get("loan_details", True):
        doc.add_heading('Loan Details:', level=2)
        loan_fields = template.get("loan_details_fields", {})
        
        if loan_fields.get("loan_id", True):
            doc.add_paragraph(f'Loan ID: {loan_data["loan_id"]}')
        if loan_fields.get("principal", True):
            doc.add_paragraph(f'Principal Amount: R {loan_data["principal"]:.2f}')
        if loan_fields.get("balance", True):
            doc.add_paragraph(f'Current Balance: R {loan_data["balance"]:.2f}')
        if loan_fields.get("amount_paid", True):
            doc.add_paragraph(f'Amount Paid: R {loan_data["amount_paid"]:.2f}')
        if loan_fields.get("due_date", True):
            doc.add_paragraph(f'Due Date: {loan_data["due_date"]}')
        if loan_fields.get("status", True):
            doc.add_paragraph(f'Status: {loan_data["status"]}')
    
    # PAYMENT HISTORY
    if include_sections.get("payment_history", True) and not payments_data.empty:
        doc.add_heading('Payment History:', level=2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        payment_fields = template.get("payment_history_fields", {})
        hdr_cells = table.rows[0].cells
        col_idx = 0
        
        if payment_fields.get("date", True):
            hdr_cells[col_idx].text = 'Date'
            col_idx += 1
        if payment_fields.get("amount", True):
            hdr_cells[col_idx].text = 'Amount'
            col_idx += 1
        if payment_fields.get("type", True):
            hdr_cells[col_idx].text = 'Type'
        
        for _, row in payments_data.iterrows():
            row_cells = table.add_row().cells
            col_idx = 0
            
            if payment_fields.get("date", True):
                row_cells[col_idx].text = str(row['date'])
                col_idx += 1
            if payment_fields.get("amount", True):
                row_cells[col_idx].text = f"R {row['amount']:.2f}"
                col_idx += 1
            if payment_fields.get("type", True):
                row_cells[col_idx].text = str(row['type'])
    
    company_footer = template.get("company_footer", "")
    if company_footer:
        doc.add_paragraph(company_footer)
    
    return doc

def run_template_editor():
    """Template customization interface"""
    st.subheader("⚙️ Invoice Template Manager")
    
    tab1, tab2, tab3 = st.tabs(["Edit Templates", "Preview", "Settings"])
    
    with tab1:
        st.write("### Create or Edit Templates")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            action = st.radio("Action:", ["Edit Existing", "Create New"])
        
        templates = get_all_templates()
        template_names = [t.get("name") for t in templates]
        
        if action == "Edit Existing":
            selected_template_name = st.selectbox("Select Template:", template_names)
            template = next((t for t in templates if t.get("name") == selected_template_name), None)
        else:
            new_template_name = st.text_input("Template Name:")
            if st.button("Create Template"):
                if new_template_name:
                    template = create_blank_template(new_template_name)
                    save_template(template)
                    st.success(f"Template '{new_template_name}' created!")
                    st.rerun()
            template = None
        
        if template:
            st.write(f"### Customizing: {template.get('name')}")
            
            # BASIC SETTINGS
            with st.expander("Basic Settings", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    template["title"] = st.text_input("Invoice Title:", template.get("title", "Invoice"))
                with col2:
                    template["company_footer"] = st.text_input("Footer Text:", template.get("company_footer", ""))
            
            # FONT SETTINGS
            with st.expander("Font Settings"):
                fonts = ["Arial", "Courier", "Helvetica", "Times"]
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    template["heading_font"] = st.selectbox("Heading Font:", fonts, 
                        index=fonts.index(template.get("heading_font", "Arial")))
                with col2:
                    template["heading_size"] = st.slider("Heading Size:", 8, 16, template.get("heading_size", 11))
                with col3:
                    template["heading_bold"] = st.checkbox("Heading Bold", template.get("heading_bold", True))
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    template["body_font"] = st.selectbox("Body Font:", fonts, 
                        index=fonts.index(template.get("body_font", "Arial")))
                with col2:
                    template["body_size"] = st.slider("Body Size:", 8, 14, template.get("body_size", 10))
                with col3:
                    template["body_bold"] = st.checkbox("Body Bold", template.get("body_bold", False))
            
            # SECTIONS
            with st.expander("Include Sections"):
                include = template.get("include_sections", {})
                include["client_info"] = st.checkbox("Client Information", include.get("client_info", True))
                include["loan_details"] = st.checkbox("Loan Details", include.get("loan_details", True))
                include["payment_history"] = st.checkbox("Payment History", include.get("payment_history", True))
                include["company_footer"] = st.checkbox("Company Footer", include.get("company_footer", True))
                template["include_sections"] = include
            
            # CLIENT INFO FIELDS
            if template.get("include_sections", {}).get("client_info", True):
                with st.expander("Client Information Fields"):
                    client_fields = template.get("client_info_fields", {})
                    col1, col2 = st.columns(2)
                    with col1:
                        client_fields["first_name"] = st.checkbox("First Name", client_fields.get("first_name", True))
                        client_fields["phone"] = st.checkbox("Phone", client_fields.get("phone", True))
                    with col2:
                        client_fields["last_name"] = st.checkbox("Last Name", client_fields.get("last_name", True))
                        client_fields["id_number"] = st.checkbox("ID Number", client_fields.get("id_number", True))
                    template["client_info_fields"] = client_fields
            
            # LOAN DETAILS FIELDS
            if template.get("include_sections", {}).get("loan_details", True):
                with st.expander("Loan Details Fields"):
                    loan_fields = template.get("loan_details_fields", {})
                    col1, col2 = st.columns(2)
                    with col1:
                        loan_fields["loan_id"] = st.checkbox("Loan ID", loan_fields.get("loan_id", True))
                        loan_fields["balance"] = st.checkbox("Current Balance", loan_fields.get("balance", True))
                        loan_fields["due_date"] = st.checkbox("Due Date", loan_fields.get("due_date", True))
                    with col2:
                        loan_fields["principal"] = st.checkbox("Principal Amount", loan_fields.get("principal", True))
                        loan_fields["amount_paid"] = st.checkbox("Amount Paid", loan_fields.get("amount_paid", True))
                        loan_fields["status"] = st.checkbox("Status", loan_fields.get("status", True))
                    template["loan_details_fields"] = loan_fields
            
            # PAYMENT HISTORY FIELDS
            if template.get("include_sections", {}).get("payment_history", True):
                with st.expander("Payment History Fields"):
                    payment_fields = template.get("payment_history_fields", {})
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        payment_fields["date"] = st.checkbox("Date", payment_fields.get("date", True))
                    with col2:
                        payment_fields["amount"] = st.checkbox("Amount", payment_fields.get("amount", True))
                    with col3:
                        payment_fields["type"] = st.checkbox("Type", payment_fields.get("type", True))
                    template["payment_history_fields"] = payment_fields
            
            # SAVE
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("💾 Save Template"):
                    save_template(template)
                    st.success(f"Template '{template.get('name')}' saved!")
                    st.rerun()
            
            with col2:
                if st.button("⭐ Set as Default"):
                    set_default_template(template.get("name"))
                    st.success(f"'{template.get('name')}' is now the default template!")
                    st.rerun()
            
            with col3:
                if template.get("name") != "Default" and st.button("🗑️ Delete Template"):
                    delete_template(template.get("name"))
                    st.success(f"Template '{template.get('name')}' deleted!")
                    st.rerun()
    
    with tab2:
        st.write("### Preview Invoice")
        st.info("Preview will be shown when you generate an invoice below.")
    
    with tab3:
        st.write("### Template Information")
        default_template = get_default_template()
        st.write(f"**Default Template:** {default_template.get('name')}")
        st.write(f"**Total Templates:** {len(templates)}")

def run(get_db):
    st.header("📄 Invoice Generator")
    
    # Sidebar for template management
    with st.sidebar:
        if st.button("⚙️ Manage Templates"):
            st.session_state.show_template_editor = True
    
    if st.session_state.get("show_template_editor", False):
        if st.button("← Back to Invoice Generator"):
            st.session_state.show_template_editor = False
            st.rerun()
        run_template_editor()
        return
    
    try:
        with get_db() as conn:
            df_loans = pd.read_sql_query("""
                SELECT l.loan_id, l.client_id, l.principal, l.balance, l.amount_paid, l.due_date, l.status,
                       c.first_name, c.last_name, c.id_number, c.phone
                FROM loans l
                JOIN clients c ON l.client_id = c.client_id
                WHERE l.status = 'Active'
            """, conn)
    except Exception as e:
        st.error(f"Database Error: {e}")
        return
    
    if df_loans.empty:
        st.warning("No active loans found to generate invoices.")
        return
    
    # Select loan
    st.subheader("Select Loan for Invoice")
    selected_loan = st.selectbox(
        "Choose a loan:",
        df_loans['loan_id'].tolist(),
        format_func=lambda x: f"Loan {x} - {df_loans[df_loans['loan_id']==x]['first_name'].iloc[0]} {df_loans[df_loans['loan_id']==x]['last_name'].iloc[0]} - R {df_loans[df_loans['loan_id']==x]['balance'].iloc[0]:.2f}"
    )
    
    # Template selection
    templates = get_all_templates()
    template_names = [t.get("name") for t in templates]
    default_idx = next((i for i, t in enumerate(templates) if t.get("is_default")), 0)
    
    selected_template_name = st.selectbox(
        "Invoice Template:",
        template_names,
        index=default_idx
    )
    
    selected_template = next((t for t in templates if t.get("name") == selected_template_name), None)
    
    if st.button("👁️ Preview Invoice"):
        loan_data = df_loans[df_loans['loan_id'] == selected_loan].iloc[0].to_dict()
        
        with get_db() as conn:
            payments_data = pd.read_sql_query("""
                SELECT date, amount, type
                FROM payment_history
                WHERE loan_id = ?
                ORDER BY date DESC
            """, conn, params=(selected_loan,))
        
        # Generate Word document preview
        doc = generate_invoice_docx(loan_data, loan_data, payments_data, selected_template)
        
        st.success("Preview generated (Word format)")
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        
        st.download_button(
            label="📥 Download Word Preview",
            data=doc_bytes.getvalue(),
            file_name=f"Invoice_Preview_{selected_loan}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    if st.button("✅ Generate Final Invoice"):
        loan_data = df_loans[df_loans['loan_id'] == selected_loan].iloc[0].to_dict()
        
        with get_db() as conn:
            payments_data = pd.read_sql_query("""
                SELECT date, amount, type
                FROM payment_history
                WHERE loan_id = ?
                ORDER BY date DESC
            """, conn, params=(selected_loan,))
        
        col1, col2 = st.columns(2)
        
        with col1:
            # PDF
            pdf = generate_invoice_pdf(loan_data, loan_data, payments_data, selected_template)
            pdf_bytes = pdf.output(dest='S').encode('latin1')
            
            st.download_button(
                label="📄 Download PDF",
                data=pdf_bytes,
                file_name=f"Invoice_{selected_loan}_{datetime.now().strftime('%Y%m%d')}.pdf",
                mime="application/pdf"
            )
        
        with col2:
            # Word (Editable)
            doc = generate_invoice_docx(loan_data, loan_data, payments_data, selected_template)
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            
            st.download_button(
                label="📝 Download Word (Editable)",
                data=doc_bytes.getvalue(),
                file_name=f"Invoice_{selected_loan}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        st.success("Invoice generated successfully!")
